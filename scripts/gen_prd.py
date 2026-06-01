"""Generate PRD .docx for Menu4U Floor Plan & Shift Manager"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.left_margin = section.right_margin = Cm(2.5)
section.top_margin  = section.bottom_margin = Cm(2.5)

# ── RTL paragraph helper ──────────────────────────────────────────────
def set_rtl(para):
    pPr = para._p.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    pPr.append(bidi)
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), 'right')
    pPr.append(jc)

def set_rtl_run(run):
    rPr = run._r.get_or_add_rPr()
    rtl = OxmlElement('w:rtl')
    rPr.append(rtl)

# ── Style helpers ─────────────────────────────────────────────────────
GOLD  = RGBColor(0xC9, 0xA8, 0x4C)
DARK  = RGBColor(0x1A, 0x12, 0x08)
GRAY  = RGBColor(0x44, 0x44, 0x44)
LGRAY = RGBColor(0x77, 0x77, 0x77)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
BLUE  = RGBColor(0x1F, 0x50, 0x9E)

def heading1(text):
    p = doc.add_paragraph()
    set_rtl(p)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = GOLD
    set_rtl_run(run)
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    return p

def heading2(text):
    p = doc.add_paragraph()
    set_rtl(p)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = DARK
    set_rtl_run(run)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    return p

def heading3(text):
    p = doc.add_paragraph()
    set_rtl(p)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = BLUE
    set_rtl_run(run)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    return p

def body(text, bold=False, color=None):
    p = doc.add_paragraph()
    set_rtl(p)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(11)
    if color:
        run.font.color.rgb = color
    set_rtl_run(run)
    p.paragraph_format.space_after = Pt(3)
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    set_rtl(p)
    run = p.add_run(text)
    run.font.size = Pt(11)
    set_rtl_run(run)
    p.paragraph_format.left_indent  = Cm(0.5 + level * 0.8)
    p.paragraph_format.space_after  = Pt(2)
    return p

def add_table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.RIGHT

    # Header row
    hdr = t.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = WHITE
        cell.paragraphs[0].runs[0].font.size = Pt(10)
        set_rtl(cell.paragraphs[0])
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), '1A1208')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:val'),  'clear')
        tcPr.append(shd)

    # Data rows
    for ri, row_data in enumerate(rows):
        row = t.rows[ri + 1]
        for ci, cell_text in enumerate(row_data):
            cell = row.cells[ci]
            cell.text = cell_text
            cell.paragraphs[0].runs[0].font.size = Pt(10)
            set_rtl(cell.paragraphs[0])
            if ri % 2 == 0:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:fill'), 'FAF5E4')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:val'),  'clear')
                tcPr.append(shd)

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)

    doc.add_paragraph()
    return t

def hr():
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pb  = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),  'single')
    bottom.set(qn('w:sz'),   '6')
    bottom.set(qn('w:space'),'1')
    bottom.set(qn('w:color'),'C9A84C')
    pb.append(bottom)
    pPr.append(pb)
    p.paragraph_format.space_after = Pt(6)

# ════════════════════════════════════════════════════════════════════
#  COVER PAGE
# ════════════════════════════════════════════════════════════════════
cover = doc.add_paragraph()
cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_rtl(cover)
r1 = cover.add_run("\nMenu4U\n")
r1.bold = True
r1.font.size = Pt(32)
r1.font.color.rgb = GOLD

cover2 = doc.add_paragraph()
cover2.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_rtl(cover2)
r2 = cover2.add_run("אפיון פונקציונלי — מערכת ניהול רצפת מסעדה ו-POS הוליסטית")
r2.bold = True
r2.font.size = Pt(16)
r2.font.color.rgb = DARK

cover3 = doc.add_paragraph()
cover3.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_rtl(cover3)
r3 = cover3.add_run(f"גרסה 1.0 · {datetime.date.today().strftime('%d.%m.%Y')}")
r3.font.size = Pt(12)
r3.font.color.rgb = LGRAY

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
#  1. CHANGE LOG
# ════════════════════════════════════════════════════════════════════
heading1("📋 יומן שינויים")
hr()
add_table(
    ["תאריך", "גרסה", "תיאור השינוי", "עודכן על ידי"],
    [
        [datetime.date.today().strftime('%d.%m.%Y'), "1.0", "יצירת מסמך ראשוני — אפיון מאוחד (PRD משתמש + הצעת מערכת)", "Claude / Menu4U"],
        ["", "—", "— שמור לשינויים עתידיים —", ""],
    ],
    col_widths=[3, 2, 8, 4]
)

# ════════════════════════════════════════════════════════════════════
#  2. OVERVIEW
# ════════════════════════════════════════════════════════════════════
heading1("1. סקירה כללית")
hr()
body("מסמך זה מגדיר את הדרישות הפונקציונליות למערכת ניהול רצפת מסעדה ו-POS הוליסטית, הכוללת שלושה מסכים עיקריים:")
bullet("מסך מלצר + Floor Plan — ממשק עבודה יומיומי לצוות השירות")
bullet("מסך מטבח KDS — הרחבת המסך הקיים")
bullet("מסך מנהל משמרת — לוח בקרה לשליטה ורווחיות")
body("")
body("המערכת בנויה על גבי התשתית הקיימת:", bold=True)

add_table(
    ["רכיב קיים", "משמש עבור"],
    [
        ["FIRE / HOLD (OrdersClient)", "שליחת מנות למטבח"],
        ["SSE stream (/api/admin/orders/stream)", "עדכונים בזמן אמת בכל המסכים"],
        ["close-table API", "סגירת חשבון + טיפ + אמצעי תשלום"],
        ["KDS (/admin/kitchen-table)", "מסך מטבח — הרחבה בלבד"],
        ["LayoutV2 (tableLayoutJson)", "מפת שולחנות — שימוש ישיר"],
        ["/api/admin/orders/waiter", "יצירת הזמנה ממסך מלצר"],
        ["AuditLog (prisma)", "תיעוד פעולות מנהל"],
    ],
    col_widths=[8, 9]
)

# ════════════════════════════════════════════════════════════════════
#  3. SCREEN 1 — FLOOR PLAN + WAITER
# ════════════════════════════════════════════════════════════════════
heading1("2. מסך 1 — Floor Plan + מסך מלצר")
hr()
body("/admin/waiter-floor  (חדש)", bold=True)
doc.add_paragraph()

heading2("2.1 מפת שולחנות (Floor Plan)")
body("תצוגה ויזואלית של הלייאוט המלא (חדרים / מרפסות / קירות) לפי הנתונים מ-tableLayoutJson. מתעדכנת בזמן אמת דרך SSE.")
doc.add_paragraph()

heading3("סטטוסי שולחן וצבעים")
add_table(
    ["צבע", "סטטוס", "תנאי הפעלה"],
    [
        ["🟢 ירוק", "פנוי", "אין הזמנות פתוחות לשולחן"],
        ["🟠 כתום", "תפוס", "קיימת הזמנה פתוחה (CONFIRMED / PREPARING / READY)"],
        ["🔴 אדום מהבהב", "ביקש חשבון", "המלצר לחץ 'ביקש חשבון' — סימון לצוות להתכונן לפינוי"],
        ["⬜ אפור", "חסום", "שולחן לא בשירות (הגדרת מנהל)"],
    ],
    col_widths=[3.5, 4, 9.5]
)

heading3("מידע על גבי השולחן")
bullet("⏱️ טיימר ישיבה — פורמט mm:ss, מתחיל מ-FIRE הראשון, נעצר בסגירת חשבון")
bullet("מספר סועדים — מוצג מתחת לטיימר")
bullet("שם מלצר אחראי (אם שויך)")
bullet("סכום נוכחי פתוח (₪)")

heading2("2.2 לחיצה על שולחן — תרחישים")
heading3("שולחן פנוי → פתיחת הזמנה חדשה")
bullet("מספר שולחן ממולא אוטומטית (ללא הקלדה ידנית)")
bullet("הזנת מספר סועדים — מוני +/- מהירים")
bullet("בחירת מנות מהתפריט — לפי קטגוריות + חיפוש")
bullet("שיוך מנה לכיסא ספציפי (כסא 1–N) לצורך פיצול עתידי")
bullet("הערה לפריט / הערה כללית לשולחן")
bullet("FIRE / HOLD לפי course (ממנגנון קיים)")
bullet("שליחה למטבח → סטטוס CONFIRMED")

heading3("שולחן תפוס → ניהול שולחן פעיל")
bullet("רשימת מנות פעילות + סטטוס כל פריט + זמן הכנה")
bullet("הוספת מנות נוספות להזמנה קיימת")
bullet("העברת שולחן — העברת כל ההזמנות + הטיימר לשולחן אחר")
bullet("איחוד שולחנות — מיזוג חשבון שולחן נוסף לחשבון אחד")
bullet("פיצול חשבון — לפי כיסאות שהוגדרו / ידנית")
bullet("'ביקש חשבון' → שינוי סטטוס לאדום מהבהב + הדפסת בון זמני")
bullet("סגירת חשבון — בחירת אמצעי תשלום + טיפ + קופון נאמנות")

heading2("2.3 טיימר ישיבה — לוגיקה")
add_table(
    ["אירוע", "פעולה על הטיימר"],
    [
        ["FIRE ראשון לשולחן", "מתחיל טיימר (t=0)"],
        ["הוספת הזמנות נוספות", "הטיימר ממשיך — לא מתאפס"],
        ["העברת שולחן", "הטיימר עובר עם ההזמנה"],
        ["'ביקש חשבון'", "הטיימר ממשיך לרוץ"],
        ["סגירת חשבון (PAID)", "הטיימר נעצר ונאפס"],
    ],
    col_widths=[7, 10]
)

# ════════════════════════════════════════════════════════════════════
#  4. SCREEN 2 — KDS
# ════════════════════════════════════════════════════════════════════
heading1("3. מסך 2 — KDS (הרחבת הקיים)")
hr()
body("/admin/kitchen-table  (הרחבה)", bold=True)
doc.add_paragraph()

heading2("שינויים על גבי המסך הקיים")
bullet("הפרדה ויזואלית בולטת: FIRE (עיגול ירוק / מיידי) vs HOLD (אפור / ממתין)")
bullet("כפתור 'מוכן להגשה' — שידור SSE למכשיר המלצר האחראי")
bullet("הצגת שיוך כיסא לכל פריט (אם הוזן מהמלצר)")
bullet("הצגת מנות שאזלו (מ-toggle של מנהל משמרת) — מוצגות עם חציה אדומה")
bullet("הערות שינוי מודגשות: 'בלי בצל', 'רוטב בצד', מידת עשייה")

# ════════════════════════════════════════════════════════════════════
#  5. SCREEN 3 — SHIFT MANAGER
# ════════════════════════════════════════════════════════════════════
heading1("4. מסך 3 — מנהל משמרת")
hr()
body("/admin/shift-manager  (חדש)", bold=True)
doc.add_paragraph()

heading2("4.1 KPIs בזמן אמת")
add_table(
    ["מדד", "חישוב", "תדירות עדכון"],
    [
        ["פדיון משמרת", "סה\"כ totalAmount של הזמנות PAID מתחילת המשמרת", "בכל סגירת שולחן"],
        ["פיצול מזומן / אשראי", "לפי paymentMethod בסגירת שולחן", "בכל סגירת שולחן"],
        ["אחוז תפוסה", "(שולחנות תפוסים / סה\"כ שולחנות) × 100", "SSE — זמן אמת"],
        ["זמן ישיבה ממוצע", "ממוצע טיימרים פעילים (דקות)", "עדכון כל דקה"],
        ["שולחנות ממתינים לתשלום", "שולחנות בסטטוס 'ביקש חשבון'", "SSE — זמן אמת"],
    ],
    col_widths=[5, 8, 4]
)

heading2("4.2 מפה מוקטנת + שיוך מלצרים")
bullet("אותה Floor Plan — read-only עם שמות מלצרים על כל שולחן")
bullet("Drag & Drop: גרירת מלצר על שולחן לשיוך")
bullet("מלצר רואה בmobile שלו רק את השולחנות המשויכים אליו (אופציה ניתנת להגדרה)")

heading2("4.3 Toggle זמינות תפריט (86 — אזלה)")
bullet("רשימת כל המנות עם מתג On/Off")
bullet("כיבוי מנה → מחשיך אותה מיידית אצל כל המלצרים ובתפריט הלקוח")
bullet("שיחזור אוטומטי של כל המנות בתחילת יום (סידול משמרת)")

heading2("4.4 פעולות מוגנות PIN מנהל")
body("כל פעולה רגישה מחייבת הזנת PIN מנהל (קוד קבוע המוגדר בהגדרות):")
bullet("ביטול מנה שכבר שודרה למטבח")
bullet("הנחה חריגה / זיכוי")
bullet("פתיחת שולחן שנסגר")
bullet("שינוי מחיר ידני")

heading3("Audit Trail — תיעוד אוטומטי")
body("כל פעולה מוגנת נרשמת ב-AuditLog עם:")
bullet("שעה ותאריך")
bullet("מספר שולחן")
bullet("שם המלצר שביצע")
bullet("סוג הפעולה + פרטים (מנה שבוטלה, סכום הנחה)")
bullet("שם המנהל שאישר ב-PIN")

heading2("4.5 התראות SLA")
add_table(
    ["סוג התראה", "תנאי", "פעולה"],
    [
        ["שולחן תקוע", "תפוס מעל X דקות ללא הזמנה נוספת", "התראה צהובה / כתומה על השולחן"],
        ["מנה מאחרת", "פריט במטבח מעל Y דקות ללא 'מוכן'", "התראה אדומה על כרטיסיית KDS"],
        ["חשבון ממתין", "שולחן ב'ביקש חשבון' מעל Z דקות", "הבהוב + קול"],
    ],
    col_widths=[5, 7, 5]
)
body("X, Y, Z — ניתנים להגדרה על ידי המנהל בהגדרות המסעדה.", color=LGRAY)

heading2("4.6 רשימת המתנה (Waitlist)")
bullet("רישום שם לקוח + מספר סועדים + שעת הגעה")
bullet("המערכת מציעה שולחן שיתפנה לפי הטיימרים הפעילים")
bullet("שיוך → שולחן עובר לסטטוס 'שמור' (אפור מוצהב) עד הגעת הלקוח")
bullet("כשהשולחן נסגר — הרשימה מקבלת עדכון אוטומטי")

heading2("4.7 סגירת משמרת")
bullet("כפתור 'סגור משמרת' — מחייב אישור מנהל")
bullet("מסמך סיכום: הכנסות, שולחנות שירותו, מספר סועדים, ממוצע לשולחן, ביטולים, סה\"כ טיפים")
bullet("ייצוא PDF / Excel")
bullet("תיעוד ב-AuditLog")

# ════════════════════════════════════════════════════════════════════
#  6. OUT OF SCOPE
# ════════════════════════════════════════════════════════════════════
heading1("5. מחוץ לטווח (שלב זה)")
hr()
bullet("Push notifications אמיתיות (נשתמש ב-SSE)")
bullet("אינטגרציה לפיסקלית / מדפסת קבלות")
bullet("שיוך מושב בממשק לקוח הסופי")
bullet("ניהול הזמנות מראש (Reservations)")

# ════════════════════════════════════════════════════════════════════
#  7. OPEN QUESTIONS
# ════════════════════════════════════════════════════════════════════
heading1("6. שאלות פתוחות לאישור")
hr()
add_table(
    ["#", "שאלה", "סטטוס"],
    [
        ["1", "מסך מלצר — panel בתוך ה-floor או ניווט לדף נפרד?", "ממתין לאישור"],
        ["2", "שיוך מלצר לשולחן — חובה בשלב זה או אופציונלי?", "ממתין לאישור"],
        ["3", "Waitlist — הצעה אוטומטית לפי טיימרים או רשימה ידנית?", "ממתין לאישור"],
        ["4", "PIN מנהל — קוד קבוע בהגדרות או מנגנון OTP?", "ממתין לאישור"],
        ["5", "86 Toggle — מתאפס אוטומטי בתחילת יום?", "ממתין לאישור"],
    ],
    col_widths=[1, 12, 4]
)

# ════════════════════════════════════════════════════════════════════
#  SAVE
# ════════════════════════════════════════════════════════════════════
out_path = "/home/user/Menu4U/Menu4U_PRD_Floor_Waiter_Shift.docx"
doc.save(out_path)
print(f"Saved: {out_path}")
